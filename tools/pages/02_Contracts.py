from __future__ import annotations

import sys

# Contracts Viewer (Streamlit page)
# Browse individual contracts under dataset/documents/contracts/<template>/ and view:
# - Paragraphs of the contract
# - Clauses (text + clause + confidence) from classification results
# - Full text
# - Datapoints extracted
# Use the sidebar to select a template and model, and navigate with "Next contract".
from pathlib import Path
from typing import Any, Optional

import pandas as pd
import streamlit as st


def get_repo_root() -> Path:
    # Find the 'dataset' directory in parents, then return its parent
    here = Path(__file__).resolve()
    dataset_dir: Optional[Path] = None
    for p in here.parents:
        if p.name == "dataset":
            dataset_dir = p
            break
    return dataset_dir.parent if dataset_dir else here.parents[2]


def list_templates() -> list[str]:
    repo_root = get_repo_root()
    # Prefer keys defined in contract_types JSONs; fallback to folder names under documents/contracts
    ct_dir = repo_root / "dataset" / "contract_types"
    keys: set[str] = set()
    try:
        for p in sorted(ct_dir.glob("*.json")):
            keys.add(p.stem)
    except Exception:
        pass
    docs_dir = repo_root / "dataset" / "documents" / "contracts"
    if docs_dir.exists():
        for p in docs_dir.iterdir():
            if p.is_dir():
                keys.add(p.name)
    return sorted(keys)


def list_models_for_template(template_key: str) -> list[str]:
    repo_root = get_repo_root()
    models: set[str] = set()
    for category in ("clauses", "datapoints"):
        base = repo_root / "dataset" / "output" / category / template_key
        if base.exists():
            for d in base.iterdir():
                if d.is_dir():
                    models.add(d.name)
    return sorted(models)


def find_contract_files(template_key: str) -> list[Path]:
    repo_root = get_repo_root()
    base = repo_root / "dataset" / "documents" / "contracts" / template_key
    files = []
    if base.exists():
        files.extend(sorted(base.glob("*.txt")))
        files.extend(sorted(base.glob("*.md")))
    return files


def load_template_dict(template_key: str) -> dict[str, Any]:
    # Add dataset/ to sys.path so we can import utilities
    dataset_dir = get_repo_root() / "dataset"
    if str(dataset_dir) not in sys.path:
        sys.path.insert(0, str(dataset_dir))
    from utilities import load_template  # type: ignore

    return load_template(template_key)


def read_text_best_effort(path: Path) -> str:
    encodings = [
        "utf-8",
        "utf-8-sig",
        "cp1252",
        "latin-1",
        "utf-16",
        "utf-16-le",
        "utf-16-be",
    ]
    for enc in encodings:
        try:
            return path.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def read_docx_text(path: Path) -> str:
    """Read text from a .docx file."""
    try:
        import docx  # type: ignore

        doc = docx.Document(str(path))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        return "\n\n".join(paragraphs)
    except Exception as e:
        raise ValueError(f"Failed to read .docx file: {e}") from e


def load_classification_csv(
    template_key: str, model_name: str, stem: str
) -> Optional[pd.DataFrame]:
    repo_root = get_repo_root()
    path = repo_root / "dataset" / "output" / "clauses" / template_key / model_name / f"{stem}.csv"
    if not path.exists():
        return None
    try:
        df = pd.read_csv(path).fillna("")
        return df
    except Exception:
        return None


def load_datapoints_csv(template_key: str, model_name: str, stem: str) -> Optional[pd.DataFrame]:
    repo_root = get_repo_root()
    path = (
        repo_root / "dataset" / "output" / "datapoints" / template_key / model_name / f"{stem}.csv"
    )
    if not path.exists():
        return None
    try:
        df = pd.read_csv(path).fillna("")
        return df
    except Exception:
        return None


def load_guidelines_csv(template_key: str, model_name: str, stem: str) -> Optional[pd.DataFrame]:
    repo_root = get_repo_root()
    path = (
        repo_root / "dataset" / "output" / "guidelines" / template_key / model_name / f"{stem}.csv"
    )
    if not path.exists():
        return None
    try:
        df = pd.read_csv(path).fillna("")
        return df
    except Exception:
        return None


def main() -> None:
    st.set_page_config(page_title="Contracts Viewer", layout="wide")
    st.title("Contracts")

    # Sidebar selectors
    st.sidebar.header("Selection")
    if "process" not in st.session_state:
        st.session_state.process = "Processing"
    process = st.sidebar.selectbox(
        "Process",
        ["Review", "Processing"],
        index=["Review", "Processing"].index(st.session_state.get("process", "Review")),
    )
    if process != st.session_state.get("process"):
        st.session_state.process = process
        st.session_state.contract_idx = 0

    templates = list_templates()
    if not templates:
        st.info("No contract types found.")
        return
    if "contracts_template" not in st.session_state:
        st.session_state.contracts_template = "ISDA"
    template_key = st.sidebar.selectbox(
        "Contract type",
        templates,
        index=templates.index(st.session_state.get("contracts_template", templates[0])),
    )
    if template_key != st.session_state.get("contracts_template"):
        st.session_state.contracts_template = template_key
        st.session_state.contract_idx = 0

    models = list_models_for_template(template_key)
    if not models:
        st.sidebar.info("No models found under output for this template.")
        model_name = None
    else:
        if "contracts_model" not in st.session_state:
            st.session_state.contracts_model = models[0]
        model_name = st.sidebar.selectbox(
            "Model", models, index=models.index(st.session_state.get("contracts_model", models[0]))
        )
        if model_name != st.session_state.get("contracts_model"):
            st.session_state.contracts_model = model_name
            st.session_state.contract_idx = 0

    # Sidebar: Create new contract uploader
    with st.sidebar.expander("Create contract", expanded=False):
        sb_uploaded = st.file_uploader(
            "Upload .txt", type=["txt"], key="sidebar_create_contract_uploader"
        )
        effective_model_sb = model_name or "gpt-4.1-mini"
        sb_disabled = sb_uploaded is None
        if st.button("Process and create", disabled=sb_disabled, key="sidebar_create_button"):
            try:
                repo_root = get_repo_root()
                txt_bytes = sb_uploaded.getvalue()  # type: ignore[union-attr]
                try:
                    text_content = txt_bytes.decode("utf-8", errors="ignore")
                except Exception:
                    text_content = txt_bytes.decode("latin-1", errors="ignore")

                # Save a copy under organizer input files
                org_in_dir = repo_root / "dataset" / "documents" / "organizer" / "files"
                org_in_dir.mkdir(parents=True, exist_ok=True)
                stem = Path(sb_uploaded.name).stem  # type: ignore[arg-type]
                (org_in_dir / f"{stem}.txt").write_text(text_content, encoding="utf-8")

                # Step 1: Organizer
                st.info("Step 1/5: Running organizer …")
                src_dir = repo_root / "src"
                if str(src_dir) not in sys.path:
                    sys.path.insert(0, str(src_dir))
                tools_dir = repo_root / "tools"
                if str(tools_dir) not in sys.path:
                    sys.path.insert(0, str(tools_dir))
                from contract_ai_core.organizer import ContractOrganizer, ContractOrganizerConfig
                from contract_ai_core.schema import Paragraph
                from utilities import load_lookup_values  # type: ignore

                temperature = 1.0 if "gpt-5" in str(effective_model_sb) else 0.2
                organizer = ContractOrganizer(
                    ContractOrganizerConfig(
                        provider="openai",
                        model=str(effective_model_sb),
                        temperature=temperature,
                        max_tokens=8000,
                        lookup_contract_types=load_lookup_values("CONTRACT_TYPE"),
                        lookup_version_types=load_lookup_values("VERSION_TYPE"),
                        lookup_statuses=load_lookup_values("STATUS"),
                    )
                )
                paragraphs = [
                    Paragraph(text=p, index=i) for i, p in enumerate(text_content.split("\n"))
                ]
                with st.spinner("Organizing …"):
                    org_results = organizer.organize([(sb_uploaded.name, paragraphs)])  # type: ignore[arg-type]
                if not org_results:
                    st.error("Organizer returned no result.")
                    st.stop()
                org_res = org_results[0]
                # Persist organizer JSON
                import json as _json

                out_org_dir = (
                    repo_root / "dataset" / "output" / "organizer" / str(effective_model_sb)
                )
                out_org_dir.mkdir(parents=True, exist_ok=True)
                out_org_path = out_org_dir / f"{stem}.json"
                try:
                    data = org_res.model_dump()  # type: ignore[attr-defined]
                except Exception:

                    def _pack(field):
                        return {
                            "value": getattr(field, "value", None),
                            "confidence": getattr(field, "confidence", None),
                            "explanation": getattr(field, "explanation", None),
                        }

                    data = {
                        "filename": getattr(org_res, "filename", sb_uploaded.name),
                        "contract_type": _pack(getattr(org_res, "contract_type", None)),
                        "contract_type_version": _pack(
                            getattr(org_res, "contract_type_version", None)
                        ),
                        "contract_date": _pack(getattr(org_res, "contract_date", None)),
                        "amendment_date": _pack(getattr(org_res, "amendment_date", None)),
                        "amendment_number": _pack(getattr(org_res, "amendment_number", None)),
                        "version_type": _pack(getattr(org_res, "version_type", None)),
                        "status": _pack(getattr(org_res, "status", None)),
                        "party_name_1": _pack(getattr(org_res, "party_name_1", None)),
                        "party_role_1": _pack(getattr(org_res, "party_role_1", None)),
                        "party_name_2": _pack(getattr(org_res, "party_name_2", None)),
                        "party_role_2": _pack(getattr(org_res, "party_role_2", None)),
                        "document_quality": _pack(getattr(org_res, "document_quality", None)),
                    }
                out_org_path.write_text(
                    _json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8"
                )
                # Also store in gold for use by the Contract tab
                gold_dir = repo_root / "dataset" / "gold" / "organizer"
                gold_dir.mkdir(parents=True, exist_ok=True)
                gold_path = gold_dir / f"{stem}.json"
                gold_path.write_text(
                    _json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8"
                )
                # Also store in gold for use by the Contract tab
                gold_dir = repo_root / "dataset" / "gold" / "organizer"
                gold_dir.mkdir(parents=True, exist_ok=True)
                gold_path = gold_dir / f"{stem}.json"
                gold_path.write_text(
                    _json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8"
                )
                contract_type_val = (
                    (data.get("contract_type") or {}).get("value")
                    if isinstance(data, dict)
                    else None
                )
                if not isinstance(contract_type_val, str) or not contract_type_val.strip():
                    st.error("Organizer did not return a contract_type.")
                    st.stop()
                contract_type_val = contract_type_val.strip()

                # Step 2: Filter and save text
                st.info("Step 2/5: Running filter and saving document …")
                from contract_ai_core import (
                    ContractTypeTemplate,
                    DocumentFilter,
                    DocumentFilterConfig,
                )
                from utilities import load_template  # type: ignore

                try:
                    tmpl_dict = load_template(contract_type_val)
                    template_ct = ContractTypeTemplate.model_validate(tmpl_dict)
                except Exception as e:
                    st.warning(f"Template load failed for {contract_type_val}: {e}")
                    template_ct = None
                contracts_dir = (
                    repo_root / "dataset" / "documents" / "contracts" / contract_type_val
                )
                contracts_dir.mkdir(parents=True, exist_ok=True)
                save_path = contracts_dir / f"{stem}.txt"
                save_path.write_text(text_content, encoding="utf-8")
                if template_ct is not None:
                    doc_filter = DocumentFilter(
                        DocumentFilterConfig(provider="openai", model=str(effective_model_sb))
                    )
                    spans = doc_filter.locate_template_scopes(
                        document_text=text_content, template=template_ct
                    )
                    out_filter_dir = (
                        repo_root
                        / "dataset"
                        / "output"
                        / "filter"
                        / contract_type_val
                        / str(effective_model_sb)
                    )
                    out_filter_dir.mkdir(parents=True, exist_ok=True)
                    out_filter_path = out_filter_dir / f"{stem}.json"
                    result = {
                        "filename": f"{stem}.txt",
                        "contract_type": contract_type_val,
                        "scopes": [
                            {
                                "name": name,
                                "start_line": span[0],
                                "end_line": span[1],
                                "confidence": span[2],
                                "explanation": span[3],
                            }
                            for name, span in spans.items()
                        ],
                    }
                    out_filter_path.write_text(
                        _json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8"
                    )

                # Steps 3-5
                from contract_ai_core.schema import (
                    ClassifiedParagraph,
                    DocumentClassification,
                )
                from contract_ai_core.utilities import text_to_paragraphs  # type: ignore

                paras = [
                    Paragraph(index=i, text=pp.text)
                    for i, pp in enumerate(text_to_paragraphs(text_content))
                ]

                # Step 3: Clauses
                st.info("Step 3/5: Classifying clauses …")
                from contract_ai_core.classifier import ClauseClassifier, ClauseClassifierConfig

                classifier = ClauseClassifier(ClauseClassifierConfig(model=str(effective_model_sb)))
                with st.spinner("Classifying …"):
                    doc_cls = classifier.classify_paragraphs(
                        paragraphs=paras,
                        template=template_ct or ContractTypeTemplate.model_validate(tmpl_dict),
                    )  # type: ignore[arg-type]
                rows_cls: list[dict[str, object]] = []
                for cp in doc_cls.paragraphs:
                    try:
                        conf_pct = (
                            None
                            if cp.confidence is None
                            else int(round(float(cp.confidence) * 100.0))
                        )
                    except Exception:
                        conf_pct = None
                    rows_cls.append(
                        {
                            "index": cp.paragraph.index,
                            "text": cp.paragraph.text,
                            "clause_key": cp.clause_key or "",
                            "confidence": conf_pct if conf_pct is not None else "",
                        }
                    )
                out_cls_dir = (
                    repo_root
                    / "dataset"
                    / "output"
                    / "clauses"
                    / contract_type_val
                    / str(effective_model_sb)
                )
                out_cls_dir.mkdir(parents=True, exist_ok=True)
                out_cls_path = out_cls_dir / f"{stem}.csv"
                pd.DataFrame(rows_cls).to_csv(out_cls_path, index=False)

                # Step 4: Datapoints
                st.info("Step 4/5: Extracting datapoints …")
                from contract_ai_core.extractor import DatapointExtractor, DatapointExtractorConfig

                extraction = DatapointExtractor(
                    DatapointExtractorConfig(model=str(effective_model_sb))
                ).extract(
                    paragraphs=paras,
                    template=template_ct or ContractTypeTemplate.model_validate(tmpl_dict),  # type: ignore[arg-type]
                    classified_clauses=doc_cls,
                )
                dp_title = {
                    str(dp.key): (dp.title or str(dp.key))
                    for dp in (
                        template_ct or ContractTypeTemplate.model_validate(tmpl_dict)
                    ).datapoints
                }  # type: ignore[arg-type]
                rows_dp: list[dict[str, object]] = []
                for item in extraction.datapoints:
                    try:
                        conf_pct = (
                            None
                            if item.confidence is None
                            else int(round(float(item.confidence) * 100.0))
                        )
                    except Exception:
                        conf_pct = None
                    rows_dp.append(
                        {
                            "key": item.key,
                            "title": dp_title.get(str(item.key), str(item.key)),
                            "value": item.value,
                            "confidence": conf_pct if conf_pct is not None else "",
                            "explanation": item.explanation or "",
                            "evidence": list(item.evidence_paragraph_indices)
                            if item.evidence_paragraph_indices
                            else [],
                        }
                    )
                out_dp_dir = (
                    repo_root
                    / "dataset"
                    / "output"
                    / "datapoints"
                    / contract_type_val
                    / str(effective_model_sb)
                )
                out_dp_dir.mkdir(parents=True, exist_ok=True)
                out_dp_path = out_dp_dir / f"{stem}.csv"
                pd.DataFrame(rows_dp).to_csv(out_dp_path, index=False)

                # Step 5: Guidelines
                st.info("Step 5/5: Reviewing guidelines …")
                from contract_ai_core.reviewer import GuidelineReviewer, GuidelineReviewerConfig

                reviewer = GuidelineReviewer(GuidelineReviewerConfig(model=str(effective_model_sb)))
                reviewed = reviewer.review(
                    paragraphs=paras,
                    template=template_ct or ContractTypeTemplate.model_validate(tmpl_dict),
                    classified_clauses=doc_cls,
                )  # type: ignore[arg-type]
                key_to_text = {
                    g.key: g.guideline
                    for g in (
                        template_ct or ContractTypeTemplate.model_validate(tmpl_dict)
                    ).guidelines
                }  # type: ignore[arg-type]
                rows_gl: list[dict[str, object]] = []
                for r in reviewed:
                    try:
                        conf_pct = (
                            None
                            if r.confidence is None
                            else int(round(float(r.confidence) * 100.0))
                        )
                    except Exception:
                        conf_pct = None
                    rows_gl.append(
                        {
                            "key": r.key,
                            "guideline": key_to_text.get(r.key, r.key),
                            "guideline_matched": str(r.guideline_matched),
                            "confidence": conf_pct if conf_pct is not None else "",
                            "explanation": r.explanation or "",
                            "evidence": list(r.evidence_paragraph_indices)
                            if r.evidence_paragraph_indices
                            else [],
                        }
                    )
                out_gl_dir = (
                    repo_root
                    / "dataset"
                    / "output"
                    / "guidelines"
                    / contract_type_val
                    / str(effective_model_sb)
                )
                out_gl_dir.mkdir(parents=True, exist_ok=True)
                out_gl_path = out_gl_dir / f"{stem}.csv"
                pd.DataFrame(rows_gl).to_csv(out_gl_path, index=False)

            except Exception as e:
                st.error(f"Creation failed: {e}")
            else:
                st.success(
                    "Contract created successfully. Outputs written to organizer/filter/clauses/datapoints/guidelines folders."
                )
                try:
                    st.session_state.contracts_template = contract_type_val
                except Exception:
                    pass
                if st.button("Reload page to view new contract", key="sidebar_reload_new_contract"):
                    st.rerun()

    # Contracts list
    files = find_contract_files(template_key)
    if not files:
        st.warning(f"No contract files found in dataset/documents/contracts/{template_key}.")
        return

    if "contract_idx" not in st.session_state:
        st.session_state.contract_idx = 0
    idx = int(st.session_state.contract_idx) % len(files)
    current_path = files[idx]

    # Load metadata for all files to create selection list
    import json as _json

    repo_root = get_repo_root()
    gold_dir = repo_root / "dataset" / "gold" / "organizer"

    def get_contract_display_name(file_path, file_idx):
        """Generate a display name for a contract based on its metadata."""
        stem = file_path.stem
        json_path = gold_dir / f"{stem}.json"

        display_name = f"{file_idx + 1}. {file_path.name}"
        if json_path.exists():
            try:
                metadata = _json.loads(json_path.read_text(encoding="utf-8"))
                contract_type = (metadata.get("contract_type") or {}).get("value") or ""
                contract_type_version = (metadata.get("contract_type_version") or {}).get(
                    "value"
                ) or ""
                contract_date = (metadata.get("contract_date") or {}).get("value") or ""
                party_name_1 = (metadata.get("party_name_1") or {}).get("value") or ""
                party_name_2 = (metadata.get("party_name_2") or {}).get("value") or ""
                party_name_1_short = " ".join(party_name_1.split()[:2]) if party_name_1 else ""
                party_name_2_short = " ".join(party_name_2.split()[:2]) if party_name_2 else ""
                display_name = f"{file_idx + 1}. {contract_type} {contract_type_version} - {contract_date} - {party_name_1_short}/{party_name_2_short}"
            except Exception:
                pass  # Fall back to default display name
        return display_name

    # Generate display names for all contracts
    contract_options = [get_contract_display_name(f, i) for i, f in enumerate(files)]

    # Selectbox for contract selection
    st.selectbox(
        "Select contract:",
        options=range(len(files)),
        index=idx,
        format_func=lambda i: contract_options[i],
        key="contract_selector",
        on_change=lambda: setattr(
            st.session_state, "contract_idx", st.session_state.contract_selector
        ),
    )

    col_prev, col_next = st.columns([2, 8])
    with col_prev:
        if st.button("◀ Previous contract"):
            st.session_state.contract_idx = (int(st.session_state.contract_idx) - 1) % len(files)
            st.rerun()
    with col_next:
        if st.button("Next contract ▶"):
            st.session_state.contract_idx = (int(st.session_state.contract_idx) + 1) % len(files)
            st.rerun()

    # Tabs (prepend Contract tab)
    if process == "Processing":
        (
            tab_contract,
            tab_text,
            tab_processing,
            tab_datapoints,
            tab_clauses,
            tab_agent,
            tab_compare,
        ) = st.tabs(
            [
                "Contract",
                "Text",
                "Processing",
                "Datapoints",
                "Clauses",
                "Ask a question",
                "Compare documents",
            ]
        )
        tab_review = st.empty()
        tab_guidelines = st.empty()
    else:
        tab_contract, tab_text, tab_review, tab_clauses, tab_guidelines, tab_agent, tab_compare = (
            st.tabs(
                [
                    "Contract",
                    "Text",
                    "Review",
                    "Clauses",
                    "Guidelines",
                    "Ask a question",
                    "Compare documents",
                ]
            )
        )
        tab_datapoints = st.empty()
        tab_processing = st.empty()

    # Contract tab: show PNG preview and gold organizer metadata
    with tab_contract:
        repo_root = get_repo_root()
        import json as _json

        pngs_dir = repo_root / "dataset" / "documents" / "organizer" / "pngs"
        gold_dir = repo_root / "dataset" / "gold" / "organizer"
        stem = current_path.stem

        col_left, col_right = st.columns([1, 1])
        with col_left:
            png_candidates = [
                pngs_dir / f"{stem}.png",
                pngs_dir / f"{current_path.name}.png",
            ]
            shown = False
            for pc in png_candidates:
                if pc.exists():
                    st.image(str(pc))
                    shown = True
                    break
            if not shown:
                st.caption("No PNG preview found for this document.")

        with col_right:
            st.subheader("Metadata")
            json_path = gold_dir / f"{stem}.json"
            if not json_path.exists():
                st.info("Gold organizer JSON not found for this document.")
            else:
                try:
                    data = _json.loads(json_path.read_text(encoding="utf-8"))
                except Exception as e:
                    st.error(f"Failed to read gold JSON: {e}")
                    data = {}

                def _fmt_pct(val: float | None) -> str:
                    try:
                        return f"{int(round(float(val) * 100))}%" if val is not None else ""
                    except Exception:
                        return ""

                def _show_field(title: str, key: str) -> None:
                    node = data.get(key) or {}
                    node = node if isinstance(node, dict) else {}
                    value = node.get("value") or ""
                    conf = node.get("confidence")
                    expl = node.get("explanation") or ""
                    st.markdown(f"**{title}:** {value}")
                    meta = _fmt_pct(conf)
                    st.caption(
                        ((f"confidence: {meta}") if meta else "") + (f" — {expl}" if expl else "")
                    )

                _show_field("Contract type", "contract_type")
                _show_field("Contract type version", "contract_type_version")
                _show_field("Contract date", "contract_date")
                _show_field("Amendment date", "amendment_date")
                _show_field("Amendment number", "amendment_number")
                _show_field("Version type", "version_type")
                _show_field("Status", "status")
                _show_field("Party name 1", "party_name_1")
                _show_field("Party role 1", "party_role_1")
                _show_field("Party name 2", "party_name_2")
                _show_field("Party role 2", "party_role_2")
                _show_field("Document quality", "document_quality")

    # Text tab: display full contract text from dataset/documents/contracts/<template_key>/<stem>.txt
    with tab_text:
        repo_root = get_repo_root()
        contracts_dir = repo_root / "dataset" / "documents" / "contracts" / template_key
        txt_path = contracts_dir / f"{current_path.stem}.txt"
        fallback_path = current_path
        if txt_path.exists():
            try:
                text = read_text_best_effort(txt_path)
            except Exception as e:
                st.error(f"Failed to read {txt_path.name}: {e}")
                text = ""
        else:
            try:
                text = read_text_best_effort(fallback_path)
            except Exception as e:
                st.error(f"Failed to read {fallback_path.name}: {e}")
                text = ""
        st.text_area("Full text", value=text, height=700)

    with tab_datapoints:
        if process != "Processing":
            # st.info("Switch to Processing to view datapoints.")
            pass
        elif not model_name:
            st.info("Select a model to view datapoints.")
        else:
            # Reprocess button (extract datapoints again for this file)
            col_btn, col_sp = st.columns([1, 9])
            with col_btn:
                if st.button("Reprocess datapoints"):
                    try:
                        # Lazy imports to avoid heavy deps when not used
                        repo_root = get_repo_root()
                        src_dir = repo_root / "src"
                        if str(src_dir) not in sys.path:
                            sys.path.insert(0, str(src_dir))
                        from contract_ai_core.extractor import (
                            DatapointExtractor,
                            DatapointExtractorConfig,
                        )
                        from contract_ai_core.schema import (
                            ClassifiedParagraph,
                            ContractTypeTemplate,
                            DocumentClassification,
                            Paragraph,
                        )
                        from contract_ai_core.utilities import text_to_paragraphs  # type: ignore

                        # Load template
                        try:
                            tmpl = load_template_dict(template_key)
                            contract_type = ContractTypeTemplate.model_validate(tmpl)
                        except Exception as e:
                            st.error(f"Template load failed: {e}")
                            contract_type = None

                        # Load text and split to paragraphs
                        text = read_text_best_effort(current_path)
                        paragraphs = [
                            Paragraph(index=i, text=pp.text)
                            for i, pp in enumerate(text_to_paragraphs(text))
                        ]

                        # Optional: load classification to scope extractions
                        df_cls_for_dp = load_classification_csv(
                            template_key, model_name, current_path.stem
                        )
                        classification: DocumentClassification | None
                        if df_cls_for_dp is not None and not df_cls_for_dp.empty:
                            cls_pars: list[ClassifiedParagraph] = []
                            clause_to_pars: dict[str, list[int]] = {}
                            for _, row in df_cls_for_dp.iterrows():
                                try:
                                    idx_i = int(str(row.get("index", "")).strip())
                                except Exception:
                                    continue
                                txt = str(row.get("text", "")).strip()
                                ck = str(row.get("clause_key", "")).strip() or None
                                cls_pars.append(
                                    ClassifiedParagraph(
                                        paragraph=Paragraph(index=idx_i, text=txt),
                                        clause_key=ck,
                                        confidence=None,
                                    )
                                )
                                if ck:
                                    clause_to_pars.setdefault(ck, []).append(idx_i)
                            classification = DocumentClassification(
                                paragraphs=cls_pars,
                                clause_to_paragraphs=clause_to_pars or None,
                            )
                        else:
                            classification = None

                        # Run extractor
                        if not contract_type:
                            st.stop()
                        extractor = DatapointExtractor(
                            DatapointExtractorConfig(model=str(model_name))
                        )
                        with st.spinner("Re-extracting datapoints…"):
                            result = extractor.extract(
                                paragraphs=paragraphs,
                                template=contract_type,
                                classified_clauses=classification,
                            )

                        # Build CSV rows (value, confidence in %, explanation, evidence)
                        # Map datapoint key -> title from template for readability
                        dp_title: dict[str, str] = {
                            str(dp.key): (dp.title or str(dp.key))
                            for dp in contract_type.datapoints
                        }
                        rows: list[dict[str, object]] = []
                        for item in result.datapoints:
                            conf_pct: int | None
                            try:
                                conf_pct = (
                                    None
                                    if item.confidence is None
                                    else int(round(float(item.confidence) * 100.0))
                                )
                            except Exception:
                                conf_pct = None
                            rows.append(
                                {
                                    "key": item.key,
                                    "title": dp_title.get(str(item.key), str(item.key)),
                                    "value": item.value,
                                    "confidence": conf_pct if conf_pct is not None else "",
                                    "explanation": item.explanation or "",
                                    "evidence": (
                                        list(item.evidence_paragraph_indices)
                                        if item.evidence_paragraph_indices
                                        else []
                                    ),
                                }
                            )

                        import pandas as _pd

                        out_df = _pd.DataFrame(rows)
                        out_dir = (
                            repo_root
                            / "dataset"
                            / "output"
                            / "datapoints"
                            / template_key
                            / str(model_name)
                        )
                        out_dir.mkdir(parents=True, exist_ok=True)
                        out_path = out_dir / f"{current_path.stem}.csv"
                        out_df.to_csv(out_path, index=False)
                    except Exception as e:
                        st.error(f"Reprocess failed: {e}")
                    else:
                        st.success("Datapoints reprocessed and saved. Reloading…")
                        st.rerun()

            df_dp_tab = load_datapoints_csv(template_key, model_name, current_path.stem)
            if df_dp_tab is None or df_dp_tab.empty:
                st.info("No datapoints results found for this contract.")
            else:
                # Render datapoints one by one using markdown
                dfv = df_dp_tab.copy()

                # Load template to detect structure types and element titles
                try:
                    tmpl = load_template_dict(template_key)
                except Exception:
                    tmpl = {}
                dp_defs = (tmpl.get("datapoints") or []) if isinstance(tmpl, dict) else []
                struct_defs = (tmpl.get("structures") or []) if isinstance(tmpl, dict) else []
                enum_defs = (tmpl.get("enums") or []) if isinstance(tmpl, dict) else []

                # Build enum lookup: enum_key -> code -> description
                enum_map: dict[str, dict[str, str]] = {}
                try:
                    for e in enum_defs:
                        ekey = str(e.get("key", "")).strip()
                        if not ekey:
                            continue
                        code_to_desc: dict[str, str] = {}
                        for opt in e.get("options") or []:
                            code = str(opt.get("code", "")).strip()
                            desc = str(opt.get("description", "")).strip()
                            if code:
                                code_to_desc[code] = desc or code
                        enum_map[ekey] = code_to_desc
                except Exception:
                    enum_map = {}

                def parse_structure_type(data_type: Any | None) -> tuple[str, str | None]:
                    s = str(data_type or "").strip().lower()
                    if s.startswith("list[") and "object:" in s:
                        inside = s[s.find("[") + 1 : s.rfind("]")]
                        if inside.startswith("object:"):
                            key = inside.split(":", 1)[1].strip().strip("[]")
                            return "list_object", key
                    if s.startswith("object:"):
                        return "object", s.split(":", 1)[1].strip().strip("[]")
                    return "simple", None

                dp_key_to_struct: dict[str, tuple[str, str]] = {}
                struct_key_to_el_title: dict[str, dict[str, str]] = {}
                struct_key_to_el_meta: dict[str, dict[str, tuple[str, str | None]]] = {}
                dp_key_to_enum_meta: dict[str, tuple[str, str | None]] = {}
                try:
                    for sd in struct_defs:
                        skey = str(sd.get("structure_key", "")).strip()
                        el_map: dict[str, str] = {}
                        el_meta: dict[str, tuple[str, str | None]] = {}
                        for el in sd.get("elements") or []:
                            el_key = str(el.get("key", "")).strip()
                            el_title = str(el.get("title", el_key)).strip()
                            el_dtype = str(el.get("data_type", "")).strip()
                            el_enum = el.get("enum_key")
                            if el_key:
                                el_map[el_key] = el_title
                                el_meta[el_key] = (
                                    el_dtype,
                                    (str(el_enum).strip() if el_enum else None),
                                )
                        if skey:
                            struct_key_to_el_title[skey] = el_map
                            struct_key_to_el_meta[skey] = el_meta
                except Exception:
                    struct_key_to_el_title = {}
                    struct_key_to_el_meta = {}
                try:
                    for dp in dp_defs:
                        key = str(dp.get("key", "")).strip()
                        kind, skey = parse_structure_type(dp.get("data_type"))
                        dt = str(dp.get("data_type", "")).strip()
                        ek = dp.get("enum_key")
                        if key and skey and kind in ("object", "list_object"):
                            dp_key_to_struct[key] = (kind, skey)
                        if key:
                            dp_key_to_enum_meta[key] = (dt, (str(ek).strip() if ek else None))
                except Exception:
                    dp_key_to_struct = {}
                    dp_key_to_enum_meta = {}

                # Build clause grouping helpers (group by clause_keys from template)
                clause_key_to_title: dict[str, str] = {}
                try:
                    tmpl2 = load_template_dict(template_key)
                    clause_defs = (tmpl2.get("clauses") or []) if isinstance(tmpl2, dict) else []
                    for c in clause_defs:
                        ck = str(c.get("key", "")).strip()
                        if ck:
                            clause_key_to_title[ck] = str(c.get("title", ck)).strip()
                except Exception:
                    clause_key_to_title = {}
                dp_key_to_clause_keys: dict[str, tuple[str, ...]] = {}
                try:
                    for dp in dp_defs:
                        k = str(dp.get("key", "")).strip()
                        cks = dp.get("clause_keys") or []
                        if isinstance(cks, (list, tuple)):
                            keys = tuple(str(x).strip() for x in cks if str(x).strip())
                        else:
                            keys = tuple()
                        if k:
                            dp_key_to_clause_keys[k] = keys
                except Exception:
                    dp_key_to_clause_keys = {}

                def try_parse_json(value: Any) -> Any:
                    s = str(value).strip()
                    if not s:
                        return None
                    if not (
                        s.startswith("{")
                        or s.startswith("[")
                        or s.startswith('"{')
                        or s.startswith('"[')
                    ):
                        return None
                    try:
                        import json

                        return json.loads(s)
                    except Exception:
                        try:
                            import ast

                            return ast.literal_eval(s)
                        except Exception:
                            return None

                def format_conf_percent(val: Any) -> str:
                    try:
                        if val is None or (isinstance(val, float) and pd.isna(val)):
                            return ""
                        num = float(val)
                        if 0.0 <= num <= 1.0:
                            num *= 100.0
                        if num < 0:
                            num = 0.0
                        if num > 100:
                            num = 100.0
                        return f"{int(round(num))}%"
                    except Exception:
                        s = str(val).strip()
                        return f"{s}%" if s and not s.endswith("%") else s

                # Group rows by clause_keys heading
                try:
                    from collections import OrderedDict
                except Exception:
                    OrderedDict = dict  # type: ignore

                def group_label_for_dp_key(k: str) -> str:
                    keys = dp_key_to_clause_keys.get(k) or tuple()
                    if keys:
                        titles = [clause_key_to_title.get(ck, ck) for ck in keys]
                        return ", ".join(titles)
                    return "Unscoped"

                grouped_rows = OrderedDict()
                for _, row in dfv.iterrows():
                    k = str(row.get("key", "")).strip()
                    label = group_label_for_dp_key(k)
                    grouped_rows.setdefault(label, []).append(row)

                def translate_bool_value(raw: Any, data_type: str | None) -> str | None:
                    if data_type is None:
                        return None
                    dt = str(data_type).strip().lower()
                    if "bool" not in dt:
                        return None
                    val = raw
                    try:
                        if val is None:
                            return "No"
                        if isinstance(val, bool):
                            return "Yes" if val else "No"
                        s = str(val).strip().lower()
                        if s in ("true", "yes", "y", "1", "t"):  # truthy tokens
                            return "Yes"
                        if s in ("false", "no", "n", "0", "f", "none", "nan", ""):
                            return "No"
                        # Numeric fallback
                        num = float(s)
                        return "Yes" if num != 0.0 else "No"
                    except Exception:
                        return "No"

                def translate_enum_value(
                    raw: Any, data_type: str | None, enum_key: str | None
                ) -> str:
                    if not enum_key or not data_type:
                        return str(raw)
                    dt = str(data_type).strip().lower()
                    e_map = enum_map.get(enum_key, {})

                    def _map_one(x: Any) -> str:
                        s = str(x).strip()
                        return e_map.get(s, s)

                    if dt.startswith("list[") and "enum" in dt:
                        val = raw
                        if isinstance(val, (list, tuple)):
                            items = [_map_one(v) for v in val]
                            return ", ".join([x for x in items if str(x).strip()])
                        # Try JSON or comma-separated
                        parsed = try_parse_json(val)
                        if isinstance(parsed, (list, tuple)):
                            items = [_map_one(v) for v in parsed]
                            return ", ".join([x for x in items if str(x).strip()])
                        s = str(val)
                        parts = [p.strip() for p in s.split(",") if p.strip()]
                        return ", ".join([_map_one(p) for p in parts]) if parts else _map_one(val)
                    if "enum" in dt:
                        return _map_one(raw)
                    return str(raw)

                def render_one_dp_row(row: pd.Series) -> None:
                    title = str(row.get("title", "")).strip() or str(row.get("key", "")).strip()
                    dp_key = str(row.get("key", "")).strip()
                    value_cell = row.get("value", "")
                    conf_str = format_conf_percent(row.get("confidence", ""))
                    explanation = str(row.get("explanation", "")).strip()

                    struct_info = dp_key_to_struct.get(dp_key)
                    if struct_info:
                        kind, skey = struct_info
                        # Heading for the datapoint
                        st.markdown(f"**{title}:**\n\n `{conf_str} {explanation}`")

                        el_titles = struct_key_to_el_title.get(skey, {})
                        parsed = try_parse_json(value_cell)

                        def render_one_object(
                            obj: Any,
                            el_titles: dict[str, str] = el_titles,
                            structure_key: str = skey,
                        ) -> None:
                            if not isinstance(obj, dict):
                                return
                            keys_in_order = list(el_titles.keys()) or list(obj.keys())
                            for el_key in keys_in_order:
                                data = obj.get(el_key)
                                ttl = el_titles.get(el_key, el_key)
                                if isinstance(data, dict):
                                    val = data.get("value")
                                    # Try boolean mapping first, then enum mapping
                                    dt_meta = (
                                        struct_key_to_el_meta.get(structure_key, {}) or {}
                                    ).get(el_key)
                                    if dt_meta:
                                        el_dtype, el_enum_key = dt_meta
                                        bool_display = translate_bool_value(val, el_dtype)
                                        if bool_display is not None:
                                            display_val = bool_display
                                        else:
                                            display_val = translate_enum_value(
                                                val, el_dtype, el_enum_key
                                            )
                                    else:
                                        display_val = val
                                    c = format_conf_percent(data.get("confidence"))
                                    expl = str(data.get("explanation", "")).strip()
                                    st.markdown(
                                        f"- **{ttl}:** {display_val}\n\n`{c} {expl}`".strip()
                                    )
                                else:
                                    # Raw value: attempt boolean/enum mapping
                                    dt_meta = (
                                        struct_key_to_el_meta.get(structure_key, {}) or {}
                                    ).get(el_key)
                                    if dt_meta:
                                        el_dtype, el_enum_key = dt_meta
                                        bool_display = translate_bool_value(data, el_dtype)
                                        if bool_display is not None:
                                            display_val = bool_display
                                        else:
                                            display_val = translate_enum_value(
                                                data, el_dtype, el_enum_key
                                            )
                                    else:
                                        display_val = data
                                    st.markdown(f"- **{ttl}:** {display_val}")

                        if kind == "list_object" and isinstance(parsed, list):
                            for obj in parsed:
                                render_one_object(obj)
                                st.markdown("")
                        else:
                            render_one_object(parsed if isinstance(parsed, dict) else {})
                            st.markdown("")
                    else:
                        # Simple datapoint: map booleans to Yes/No and translate enum codes
                        dt_meta = dp_key_to_enum_meta.get(dp_key)
                        if dt_meta:
                            dt_s, enum_key = dt_meta
                            bool_display = translate_bool_value(value_cell, dt_s)
                            if bool_display is not None:
                                value_display = bool_display
                            else:
                                value_display = translate_enum_value(value_cell, dt_s, enum_key)
                        else:
                            value_display = value_cell
                        st.markdown(
                            f"**{title}**: {value_display}\n\n`{conf_str} {explanation}`".strip()
                        )
                        # st.markdown("")

                for heading, rows in grouped_rows.items():
                    st.markdown(f"### {heading}")
                    for row in rows:
                        render_one_dp_row(row)

    with tab_processing:
        if process != "Processing":
            # st.info("Switch to Processing to view classification.")
            pass
        elif not model_name:
            st.info("Select a model to view clause classification results.")
        else:
            df_cls = load_classification_csv(template_key, model_name, current_path.stem)
            if df_cls is None or df_cls.empty:
                st.info("No clause results found for this contract.")
            else:
                # Expect columns: index, clause_key, confidence (percent int), text
                text_series = df_cls.get("text", pd.Series(dtype=str))
                clause_series = df_cls.get("clause_key", pd.Series(dtype=str))
                # Map clause keys to titles from the template
                try:
                    tmpl = load_template_dict(template_key)
                    clauses = tmpl.get("clauses", []) or []
                    key_to_title = {
                        str(c.get("key")): c.get("title") or str(c.get("key")) for c in clauses
                    }
                    # Map guideline key -> priority from template
                    gl_defs = tmpl.get("guidelines", []) or []
                    gl_key_to_priority: dict[str, str] = {
                        str(g.get("key", "")).strip(): str(g.get("priority", "")).strip()
                        for g in gl_defs
                        if str(g.get("key", "")).strip()
                    }
                except Exception:
                    key_to_title = {}
                    gl_key_to_priority = {}
                clause_title_series = clause_series.map(lambda k: key_to_title.get(str(k), str(k)))
                conf_series = df_cls.get("confidence", pd.Series(dtype=object)).map(
                    lambda v: f"{int(v)}%" if pd.notna(v) and str(v).strip() != "" else ""
                )
                # Build clause cells: if same key as previous non-empty, show only (confidence) with left border
                keys_list = clause_series.tolist()
                titles_list = clause_title_series.tolist()
                conf_list = conf_series.tolist()
                repeat_flags: list[bool] = []
                prev_key: str | None = None
                for k in keys_list:
                    k_str = str(k).strip() if k is not None else ""
                    is_repeat = prev_key is not None and k_str != "" and k_str == prev_key
                    repeat_flags.append(bool(is_repeat))
                    if k_str != "":
                        prev_key = k_str

                clause_cells: list[str] = []
                text_cells: list[str] = []
                for is_rep, title, conf in zip(repeat_flags, titles_list, conf_list, strict=False):
                    conf_str = str(conf).strip()
                    if is_rep:
                        content = f'<span class="dp-conf">{conf_str}</span>' if conf_str else ""
                        clause_cells.append(f'<div class="repeat-cell">{content}</div>')
                        # Mark text cell as repeated to draw right border
                        text_cells.append('<div class="text-repeat-cell">{}</div>')
                    else:
                        title_str = str(title).strip()
                        content = (
                            f'{title_str} <span class="dp-conf">{conf_str}</span>'
                            if conf_str
                            else title_str
                        )
                        clause_cells.append(f'<div class="normal-cell">{content}</div>')
                        text_cells.append('<div class="text-normal-cell">{}</div>')

                # Merge raw text into wrappers (avoid HTML breaking)
                safe_text = (
                    text_series.astype(str)
                    .str.replace("<", "&lt;")
                    .str.replace(">", "&gt;")
                    .str.replace("\t", "")
                )
                text_cells = [
                    tpl.format(txt)
                    for tpl, txt in zip(text_cells, safe_text.tolist(), strict=False)
                ]

                # Build datapoints column aligned by first evidence paragraph index
                n_rows = len(text_cells)
                dp_cells: list[str] = [""] * n_rows
                df_dp_for_clauses = load_datapoints_csv(template_key, model_name, current_path.stem)
                if df_dp_for_clauses is not None and not df_dp_for_clauses.empty:
                    try:
                        tmpl = load_template_dict(template_key)
                    except Exception:
                        tmpl = {}
                    dp_defs = (tmpl.get("datapoints") or []) if isinstance(tmpl, dict) else []
                    struct_defs = (tmpl.get("structures") or []) if isinstance(tmpl, dict) else []

                    def parse_structure_type(data_type: Any | None) -> tuple[str, str | None]:
                        s = str(data_type or "").strip().lower()
                        if s.startswith("list[") and "object:" in s:
                            inside = s[s.find("[") + 1 : s.rfind("]")]
                            if inside.startswith("object:"):
                                key = inside.split(":", 1)[1].strip().strip("[]")
                                return "list_object", key
                        if s.startswith("object:"):
                            return "object", s.split(":", 1)[1].strip().strip("[]")
                        return "simple", None

                    dp_key_to_struct: dict[str, tuple[str, str]] = {}
                    struct_key_to_el_title: dict[str, dict[str, str]] = {}
                    try:
                        for sd in struct_defs:
                            skey = str(sd.get("structure_key", "")).strip()
                            el_map: dict[str, str] = {}
                            for el in sd.get("elements") or []:
                                el_key = str(el.get("key", "")).strip()
                                el_title = str(el.get("title", el_key)).strip()
                                if el_key:
                                    el_map[el_key] = el_title
                            if skey:
                                struct_key_to_el_title[skey] = el_map
                    except Exception:
                        struct_key_to_el_title = {}
                    try:
                        for dp in dp_defs:
                            key = str(dp.get("key", "")).strip()
                            kind, skey = parse_structure_type(dp.get("data_type"))
                            if key and skey and kind in ("object", "list_object"):
                                dp_key_to_struct[key] = (kind, skey)
                    except Exception:
                        dp_key_to_struct = {}

                    def try_parse_json(value: Any) -> Any:
                        s = str(value).strip()
                        if not s:
                            return None
                        if not (
                            s.startswith("{")
                            or s.startswith("[")
                            or s.startswith('"{')
                            or s.startswith('"[')
                        ):
                            return None
                        try:
                            import json

                            return json.loads(s)
                        except Exception:
                            try:
                                import ast

                                return ast.literal_eval(s)
                            except Exception:
                                return None

                    def parse_evidence_field(value: Any) -> list[int]:
                        try:
                            import ast
                        except Exception:
                            ast = None  # type: ignore
                        if value is None or (isinstance(value, float) and pd.isna(value)):
                            return []
                        s = str(value).strip()
                        if s == "":
                            return []
                        try:
                            if "ast" in locals() and ast is not None:
                                parsed = ast.literal_eval(s)
                                if isinstance(parsed, (list, tuple)):
                                    out: list[int] = []
                                    for x in parsed:
                                        try:
                                            out.append(int(x))
                                        except Exception:
                                            continue
                                    return out
                                if isinstance(parsed, (int, float)):
                                    return [int(parsed)]
                        except Exception:
                            pass
                        tokens = s.replace("[", "").replace("]", "").split(",")
                        out: list[int] = []
                        for tok in tokens:
                            tok = tok.strip()
                            if not tok:
                                continue
                            try:
                                out.append(int(float(tok)))
                            except Exception:
                                continue
                        return out

                    for _, r in df_dp_for_clauses.iterrows():
                        dp_key = str(r.get("key", "")).strip()
                        title = str(r.get("title", "")).strip()
                        value_str = str(r.get("value", "")).strip()
                        conf_val = r.get("confidence", "")
                        conf_str = ""
                        try:
                            if pd.notna(conf_val) and str(conf_val).strip() != "":
                                conf_str = f"{int(conf_val)}%"
                        except Exception:
                            conf_str = str(conf_val).strip()
                        explanation = str(r.get("explanation", "")).strip()

                        struct_info = dp_key_to_struct.get(dp_key)
                        if struct_info:
                            kind, skey = struct_info
                            header_line = f"{title}:"
                            parts: list[str] = [header_line]
                            meta_conf = str(conf_str).strip()
                            meta_expl = str(explanation).strip()
                            meta_parts: list[str] = []
                            if meta_conf:
                                meta_parts.append(f'<span class="dp-conf">{meta_conf}</span>')
                            if meta_expl:
                                meta_parts.append(f'<span class="dp-meta">{meta_expl}</span>')
                            if meta_parts:
                                parts.append(" ".join(meta_parts))
                            parsed = try_parse_json(r.get("value", ""))
                            el_titles = struct_key_to_el_title.get(skey, {})

                            def render_one_object(
                                obj: Any,
                                el_titles: dict[str, str] = el_titles,
                                structure_key: str = skey,
                            ) -> list[str]:
                                lines: list[str] = []
                                if isinstance(obj, dict):
                                    keys_in_order = list(el_titles.keys()) or list(obj.keys())
                                    for el_key in keys_in_order:
                                        data = obj.get(el_key)
                                        if not isinstance(data, dict):
                                            # Attempt enum translation for raw values
                                            dt_meta = (
                                                struct_key_to_el_meta.get(structure_key, {}) or {}
                                            ).get(el_key)
                                            if dt_meta:
                                                el_dtype, el_enum_key = dt_meta
                                                val = translate_enum_value(
                                                    data, el_dtype, el_enum_key
                                                )
                                            else:
                                                val = str(data)
                                            ttl = el_titles.get(el_key, el_key)
                                            sub = f"{ttl}: {val}".strip()
                                            lines.append(sub)
                                            continue
                                        val = data.get("value")
                                        c = str(data.get("confidence", "")).strip()
                                        if c and not c.endswith("%"):
                                            try:
                                                c = f"{int(float(c)*100) if float(c) <= 1.0 else int(float(c))}%"
                                            except Exception:
                                                pass
                                        expl = data.get("explanation")
                                        # Attempt enum translation for structured objects with value/confidence/explanation
                                        dt_meta = (
                                            struct_key_to_el_meta.get(structure_key, {}) or {}
                                        ).get(el_key)
                                        if dt_meta:
                                            el_dtype, el_enum_key = dt_meta
                                            display_val = translate_enum_value(
                                                val, el_dtype, el_enum_key
                                            )
                                        else:
                                            display_val = val
                                        ttl = el_titles.get(el_key, el_key)
                                        first = f"{ttl}: {display_val}".strip()
                                        second_parts: list[str] = []
                                        c_str = str(c).strip()
                                        expl_str = str(expl or "").strip()
                                        if c_str:
                                            second_parts.append(
                                                f'<span class="dp-conf">{c_str}</span>'
                                            )
                                        if expl_str:
                                            second_parts.append(
                                                f'<span class="dp-meta">{expl_str}</span>'
                                            )
                                        second = " ".join(second_parts)
                                        lines.append(first)
                                        if second:
                                            lines.append(second)
                                return lines

                            if kind == "list_object" and isinstance(parsed, list):
                                items_blocks: list[str] = []
                                for obj in parsed:
                                    block_lines = render_one_object(obj)
                                    if block_lines:
                                        items_blocks.append("<br/>".join(block_lines))
                                if items_blocks:
                                    parts.append("<br/><br/>".join(items_blocks))
                            else:
                                block_lines = render_one_object(
                                    parsed if isinstance(parsed, dict) else {}
                                )
                                if block_lines:
                                    parts.append("<br/>".join(block_lines))
                            content_html = "<br/>".join(parts)
                        else:
                            content_lines = [f"{title}: {value_str}"]
                            meta_conf = str(conf_str).strip()
                            meta_expl = str(explanation).strip()
                            meta_parts: list[str] = []
                            if meta_conf:
                                meta_parts.append(f'<span class="dp-conf">{meta_conf}</span>')
                            if meta_expl:
                                meta_parts.append(f'<span class="dp-meta">{meta_expl}</span>')
                            if meta_parts:
                                content_lines.append(" ".join(meta_parts))
                            content_html = "<br/>".join(content_lines)

                        ev = parse_evidence_field(r.get("evidence", ""))
                        idx = 0
                        if ev:
                            try:
                                idx = min(int(i) for i in ev)
                            except Exception:
                                idx = 0
                        if n_rows > 0:
                            idx = max(0, min(idx, n_rows - 1))
                        dp_cells[idx] = (
                            (dp_cells[idx] + f'<div class="dp-item">{content_html}</div>')
                            if dp_cells[idx]
                            else f'<div class="dp-item">{content_html}</div>'
                        )

                df_view = pd.DataFrame(
                    {
                        "text": pd.Series(text_cells),
                        "clause": clause_cells,
                        "datapoints": pd.Series(dp_cells),
                    }
                )
                # Render a single HTML table with fixed column widths and no borders; repeated rows show a right border on the text column
                html = df_view.to_html(index=False, escape=False)
                html = html.replace(
                    '<table border="1" class="dataframe">',
                    (
                        "<style>"
                        ".wrapped-table{width:100%;table-layout:fixed;border-collapse:collapse;}"
                        ".wrapped-table th,.wrapped-table td{border:none;padding:0.5rem;vertical-align:top;word-wrap:break-word;white-space:normal;}"
                        ".wrapped-table td .text-repeat-cell{border-right:4px solid #888;padding-right:0.5rem;}"
                        ".wrapped-table col.text-col{width:50%;}"
                        ".wrapped-table col.clause-col{width:25%;}"
                        ".wrapped-table col.dp-col{width:25%;}"
                        ".wrapped-table .dp-item{margin-bottom:0.5rem;padding-bottom:0.25rem;border-bottom:1px solid #eee;}"
                        ".wrapped-table .dp-conf{color:#2e7d32;font-weight:600;}"
                        ".wrapped-table .dp-meta{color:#666;}"
                        "</style>"
                        '<table class="wrapped-table">'
                        "<colgroup>"
                        '<col class="text-col"/>'
                        '<col class="clause-col"/>'
                        '<col class="dp-col"/>'
                        "</colgroup>"
                    ),
                )
                st.markdown(html, unsafe_allow_html=True)

    with tab_review:
        if process == "Processing":
            # st.info("Switch to Review to view guidelines alongside clauses.")
            pass
        elif not model_name:
            st.info("Select a model to review guidelines.")
        else:
            # Load classification for clause labels
            df_cls = load_classification_csv(template_key, model_name, current_path.stem)
            # Load guidelines
            df_gl = load_guidelines_csv(template_key, model_name, current_path.stem)
            if df_cls is None or df_cls.empty:
                st.info("No clause results found for this contract.")
            elif df_gl is None or df_gl.empty:
                st.info("No guidelines results found for this contract.")
            else:
                # Prepare clause and text columns (same as Clauses tab)
                text_series = df_cls.get("text", pd.Series(dtype=str))
                clause_series = df_cls.get("clause_key", pd.Series(dtype=str))
                try:
                    tmpl = load_template_dict(template_key)
                    clauses = tmpl.get("clauses", []) or []
                    key_to_title = {
                        str(c.get("key")): c.get("title") or str(c.get("key")) for c in clauses
                    }
                except Exception:
                    key_to_title = {}

                # Map guideline key -> priority from template for priority badges
                gl_key_to_priority: dict[str, str] = {}
                try:
                    tmpl_gl = load_template_dict(template_key)
                    gl_defs = tmpl_gl.get("guidelines", []) or []
                    for g in gl_defs:
                        gk = str(g.get("key", "")).strip()
                        if not gk:
                            continue
                        gl_key_to_priority[gk] = str(g.get("priority", "")).strip()
                except Exception:
                    gl_key_to_priority = {}
                clause_title_series = clause_series.map(lambda k: key_to_title.get(str(k), str(k)))
                conf_series = df_cls.get("confidence", pd.Series(dtype=object)).map(
                    lambda v: f"{int(v)}%" if pd.notna(v) and str(v).strip() != "" else ""
                )
                keys_list = clause_series.tolist()
                titles_list = clause_title_series.tolist()
                conf_list = conf_series.tolist()
                repeat_flags: list[bool] = []
                prev_key: str | None = None
                for k in keys_list:
                    k_str = str(k).strip() if k is not None else ""
                    is_repeat = prev_key is not None and k_str != "" and k_str == prev_key
                    repeat_flags.append(bool(is_repeat))
                    if k_str != "":
                        prev_key = k_str
                clause_cells: list[str] = []
                text_cells: list[str] = []
                for is_rep, title, conf in zip(repeat_flags, titles_list, conf_list, strict=False):
                    conf_str = str(conf).strip()
                    if is_rep:
                        content = f"({conf_str})" if conf_str else ""
                        clause_cells.append(f'<div class="repeat-cell">{content}</div>')
                        text_cells.append('<div class="text-repeat-cell">{}</div>')
                    else:
                        title_str = str(title).strip()
                        content = f"{title_str} ({conf_str})" if conf_str else title_str
                        clause_cells.append(f'<div class="normal-cell">{content}</div>')
                        text_cells.append('<div class="text-normal-cell">{}</div>')
                safe_text = (
                    text_series.astype(str)
                    .str.replace("<", "&lt;")
                    .str.replace(">", "&gt;")
                    .str.replace("\t", "")
                )
                text_cells = [
                    tpl.format(txt)
                    for tpl, txt in zip(text_cells, safe_text.tolist(), strict=False)
                ]

                # Build guidelines column aligned by first evidence paragraph index
                n_rows = len(text_cells)
                gl_cells: list[str] = [""] * n_rows

                def parse_evidence_field(value: Any) -> list[int]:
                    try:
                        import ast
                    except Exception:
                        ast = None  # type: ignore
                    if value is None or (isinstance(value, float) and pd.isna(value)):
                        return []
                    s = str(value).strip()
                    if s == "":
                        return []
                    try:
                        if "ast" in locals() and ast is not None:
                            parsed = ast.literal_eval(s)
                            if isinstance(parsed, (list, tuple)):
                                out: list[int] = []
                                for x in parsed:
                                    try:
                                        out.append(int(x))
                                    except Exception:
                                        continue
                                return out
                            if isinstance(parsed, (int, float)):
                                return [int(parsed)]
                    except Exception:
                        pass
                    tokens = s.replace("[", "").replace("]", "").split(",")
                    out: list[int] = []
                    for tok in tokens:
                        tok = tok.strip()
                        if not tok:
                            continue
                        try:
                            out.append(int(float(tok)))
                        except Exception:
                            continue
                    return out

                def _priority_class(p: str) -> str:
                    s = str(p or "").strip().lower()
                    if s in ("high", "material", "critical", "severe"):
                        return "sev-material"
                    if s in ("medium", "important", "major"):
                        return "sev-important"
                    if s in ("low", "minor"):
                        return "sev-minor"
                    return "sev-minor"

                for _, r in df_gl.iterrows():
                    guideline_text = str(r.get("guideline", "")).strip()
                    gkey = str(r.get("key", "")).strip()
                    prio = gl_key_to_priority.get(gkey, "")
                    matched = str(r.get("guideline_matched", "")).strip()
                    conf_val = r.get("confidence", "")
                    try:
                        conf_str = (
                            f"{int(conf_val)}%"
                            if pd.notna(conf_val) and str(conf_val).strip() != ""
                            else ""
                        )
                    except Exception:
                        conf_str = str(conf_val).strip()
                    explanation = str(r.get("explanation", "")).strip()
                    header = guideline_text
                    # Badge: Verified (green) if matched truthy; otherwise show priority chip
                    is_matched = str(matched).strip().lower() in ("true", "yes", "1")
                    if is_matched:
                        badge_html = '<span class="sev verified">Verified</span>'
                    else:
                        pr_cls = _priority_class(prio)
                        pr_label = (str(prio).strip() or "priority").capitalize()
                        badge_html = f'<span class="sev {pr_cls}">{pr_label}</span>'
                    conf_html = f'<span class="dp-conf">{conf_str}</span>' if conf_str else ""
                    expl_html = f'<span class="dp-meta">{explanation}</span>' if explanation else ""
                    meta_parts = [x for x in [conf_html, expl_html] if x]
                    meta_html = (
                        badge_html if not meta_parts else f"{badge_html} {' '.join(meta_parts)}"
                    )
                    content_html = header if not meta_html else f"{header}<br/>{meta_html}"

                    ev = parse_evidence_field(r.get("evidence", ""))
                    idx = 0
                    if ev:
                        try:
                            idx = min(int(i) for i in ev)
                        except Exception:
                            idx = 0
                    if n_rows > 0:
                        idx = max(0, min(idx, n_rows - 1))
                    gl_cells[idx] = (
                        (gl_cells[idx] + f'<div class="dp-item">{content_html}</div>')
                        if gl_cells[idx]
                        else f'<div class="dp-item">{content_html}</div>'
                    )

                df_view = pd.DataFrame(
                    {
                        "text": pd.Series(text_cells),
                        "clause": clause_cells,
                        "guidelines": pd.Series(gl_cells),
                    }
                )
                html = df_view.to_html(index=False, escape=False)
                html = html.replace(
                    '<table border="1" class="dataframe">',
                    (
                        "<style>"
                        ".wrapped-table{width:100%;table-layout:fixed;border-collapse:collapse;}"
                        ".wrapped-table th,.wrapped-table td{border:none;padding:0.5rem;vertical-align:top;word-wrap:break-word;white-space:normal;}"
                        ".wrapped-table td .text-repeat-cell{border-right:4px solid #888;padding-right:0.5rem;}"
                        ".wrapped-table col.text-col{width:50%;}"
                        ".wrapped-table col.clause-col{width:25%;}"
                        ".wrapped-table col.dp-col{width:25%;}"
                        ".wrapped-table .dp-item{margin-bottom:0.5rem;padding-bottom:0.25rem;border-bottom:1px solid #eee;}"
                        ".wrapped-table .sev{padding:2px 6px;border-radius:10px;font-size:0.85em;color:#fff;margin-right:6px;}"
                        ".wrapped-table .sev-material{background:#b71c1c;}"
                        ".wrapped-table .sev-important{background:#ef6c00;}"
                        ".wrapped-table .sev-minor{background:#616161;}"
                        ".wrapped-table .verified{background:#2e7d32;}"
                        "</style>"
                        '<table class="wrapped-table">'
                        "<colgroup>"
                        '<col class="text-col"/>'
                        '<col class="clause-col"/>'
                        '<col class="dp-col"/>'
                        "</colgroup>"
                    ),
                )
                st.markdown(html, unsafe_allow_html=True)

    with tab_clauses:
        if not model_name:
            st.info("Select a model to view clauses.")
        else:
            # Reprocess button (classify clauses again for this file)
            col_btn_cls, _ = st.columns([1, 9])
            with col_btn_cls:
                if st.button("Reprocess clauses"):
                    try:
                        # Setup imports
                        repo_root = get_repo_root()
                        src_dir = repo_root / "src"
                        if str(src_dir) not in sys.path:
                            sys.path.insert(0, str(src_dir))
                        from contract_ai_core.classifier import (
                            ClauseClassifier,
                            ClauseClassifierConfig,
                        )
                        from contract_ai_core.schema import (
                            ContractTypeTemplate,
                            Paragraph,
                        )
                        from contract_ai_core.utilities import text_to_paragraphs  # type: ignore

                        # Load template
                        try:
                            tmpl_dict = load_template_dict(template_key)
                            contract_type = ContractTypeTemplate.model_validate(tmpl_dict)
                        except Exception as e:
                            st.error(f"Template load failed: {e}")
                            st.stop()

                        # Load text and split to paragraphs
                        text = read_text_best_effort(current_path)
                        paragraphs = [
                            Paragraph(index=i, text=pp.text)
                            for i, pp in enumerate(text_to_paragraphs(text))
                        ]

                        # Run classifier
                        classifier = ClauseClassifier(ClauseClassifierConfig(model=str(model_name)))
                        with st.spinner("Reclassifying clauses…"):
                            doc_classification = classifier.classify_paragraphs(
                                paragraphs=paragraphs,
                                template=contract_type,
                                source_id=current_path.stem,
                            )

                        # Prepare CSV rows: index, text, clause_key, confidence (percent int)
                        rows: list[dict[str, object]] = []
                        for cp in doc_classification.paragraphs:
                            conf_pct: int | None
                            try:
                                conf_pct = (
                                    None
                                    if cp.confidence is None
                                    else int(round(float(cp.confidence) * 100.0))
                                )
                            except Exception:
                                conf_pct = None
                            rows.append(
                                {
                                    "index": cp.paragraph.index,
                                    "text": cp.paragraph.text,
                                    "clause_key": cp.clause_key or "",
                                    "confidence": conf_pct if conf_pct is not None else "",
                                }
                            )

                        out_dir = (
                            repo_root
                            / "dataset"
                            / "output"
                            / "clauses"
                            / template_key
                            / str(model_name)
                        )
                        out_dir.mkdir(parents=True, exist_ok=True)
                        out_path = out_dir / f"{current_path.stem}.csv"
                        pd.DataFrame(rows).to_csv(out_path, index=False)
                    except Exception as e:
                        st.error(f"Reprocess failed: {e}")
                    else:
                        st.success("Clauses reprocessed and saved. Reloading…")
                        st.rerun()

            df_cls = load_classification_csv(template_key, model_name, current_path.stem)
            if df_cls is None or df_cls.empty:
                st.info("No clause results found for this contract.")
            else:
                # Map clause keys to titles
                try:
                    tmpl = load_template_dict(template_key)
                    clauses = tmpl.get("clauses", []) or []
                    key_to_title = {
                        str(c.get("key")): c.get("title") or str(c.get("key")) for c in clauses
                    }
                except Exception:
                    key_to_title = {}

                # Group paragraphs by clause_key in appearance order
                grouped: dict[str, list[str]] = {}
                for _, row in df_cls.iterrows():
                    k = str(row.get("clause_key", "")).strip()
                    if not k:
                        continue
                    txt = str(row.get("text", "")).strip()
                    if not txt:
                        continue
                    grouped.setdefault(k, []).append(txt)

                if not grouped:
                    st.info("No classified clauses found in this document.")
                else:
                    for ck, para_list in grouped.items():
                        title = key_to_title.get(ck, ck)
                        st.markdown(f"### {title}")
                        for p in para_list:
                            st.markdown(f"- {p}")

    with tab_guidelines:
        if process == "Processing":
            # st.info("Switch to Review to view guidelines.")
            pass
        elif not model_name:
            st.info("Select a model to view guidelines.")
        else:
            # Reprocess button (review guidelines again for this file)
            col_btn_gl, _ = st.columns([1, 9])
            with col_btn_gl:
                if st.button("Reprocess guidelines"):
                    try:
                        repo_root = get_repo_root()
                        src_dir = repo_root / "src"
                        if str(src_dir) not in sys.path:
                            sys.path.insert(0, str(src_dir))
                        from contract_ai_core.reviewer import (
                            GuidelineReviewer,
                            GuidelineReviewerConfig,
                        )
                        from contract_ai_core.schema import (
                            ClassifiedParagraph,
                            ContractTypeTemplate,
                            DocumentClassification,
                            Paragraph,
                        )
                        from contract_ai_core.utilities import text_to_paragraphs  # type: ignore

                        # Load template
                        try:
                            tmpl_dict = load_template_dict(template_key)
                            contract_type = ContractTypeTemplate.model_validate(tmpl_dict)
                        except Exception as e:
                            st.error(f"Template load failed: {e}")
                            st.stop()

                        # Load text and split to paragraphs
                        text = read_text_best_effort(current_path)
                        paragraphs = [
                            Paragraph(index=i, text=pp.text)
                            for i, pp in enumerate(text_to_paragraphs(text))
                        ]

                        # Optional classification to scope guideline reviews
                        df_cls_for_gl = load_classification_csv(
                            template_key, model_name, current_path.stem
                        )
                        classification: DocumentClassification | None
                        if df_cls_for_gl is not None and not df_cls_for_gl.empty:
                            cls_pars: list[ClassifiedParagraph] = []
                            clause_to_pars: dict[str, list[int]] = {}
                            for _, row in df_cls_for_gl.iterrows():
                                try:
                                    idx_i = int(str(row.get("index", "")).strip())
                                except Exception:
                                    continue
                                txt = str(row.get("text", "")).strip()
                                ck = str(row.get("clause_key", "")).strip() or None
                                cls_pars.append(
                                    ClassifiedParagraph(
                                        paragraph=Paragraph(index=idx_i, text=txt),
                                        clause_key=ck,
                                        confidence=None,
                                    )
                                )
                                if ck:
                                    clause_to_pars.setdefault(ck, []).append(idx_i)
                            classification = DocumentClassification(
                                paragraphs=cls_pars,
                                clause_to_paragraphs=clause_to_pars or None,
                            )
                        else:
                            classification = None

                        # Run reviewer
                        reviewer = GuidelineReviewer(GuidelineReviewerConfig(model=str(model_name)))
                        with st.spinner("Re-reviewing guidelines…"):
                            reviewed = reviewer.review(
                                paragraphs=paragraphs,
                                template=contract_type,
                                classified_clauses=classification,
                            )

                        # Map key -> guideline text from template for CSV readability
                        key_to_text = {g.key: g.guideline for g in contract_type.guidelines}

                        rows: list[dict[str, object]] = []
                        for r in reviewed:
                            try:
                                conf_pct = (
                                    None
                                    if r.confidence is None
                                    else int(round(float(r.confidence) * 100.0))
                                )
                            except Exception:
                                conf_pct = None
                            rows.append(
                                {
                                    "key": r.key,
                                    "guideline": key_to_text.get(r.key, r.key),
                                    "guideline_matched": str(r.guideline_matched),
                                    "confidence": conf_pct if conf_pct is not None else "",
                                    "explanation": r.explanation or "",
                                    "evidence": (
                                        list(r.evidence_paragraph_indices)
                                        if r.evidence_paragraph_indices
                                        else []
                                    ),
                                }
                            )

                        out_dir = (
                            repo_root
                            / "dataset"
                            / "output"
                            / "guidelines"
                            / template_key
                            / str(model_name)
                        )
                        out_dir.mkdir(parents=True, exist_ok=True)
                        out_path = out_dir / f"{current_path.stem}.csv"
                        pd.DataFrame(rows).to_csv(out_path, index=False)
                    except Exception as e:
                        st.error(f"Reprocess failed: {e}")
                    else:
                        st.success("Guidelines reprocessed and saved. Reloading…")
                        st.rerun()

            df_gl = load_guidelines_csv(template_key, model_name, current_path.stem)
            if df_gl is None or df_gl.empty:
                st.info("No guidelines results found for this contract.")
            else:
                # Build clause title lookup and guideline->clause_keys mapping from template
                clause_key_to_title: dict[str, str] = {}
                gl_key_to_clause_keys: dict[str, tuple[str, ...]] = {}
                gl_key_to_priority: dict[str, str] = {}
                try:
                    tmpl2 = load_template_dict(template_key)
                    clause_defs = (tmpl2.get("clauses") or []) if isinstance(tmpl2, dict) else []
                    for c in clause_defs:
                        ck = str(c.get("key", "")).strip()
                        if ck:
                            clause_key_to_title[ck] = str(c.get("title", ck)).strip()
                    guideline_defs = (
                        (tmpl2.get("guidelines") or []) if isinstance(tmpl2, dict) else []
                    )
                    for g in guideline_defs:
                        gk = str(g.get("key", "")).strip()
                        pr = str(g.get("priority", "")).strip()
                        cks = g.get("clause_keys") or []
                        if isinstance(cks, (list, tuple)):
                            keys = tuple(str(x).strip() for x in cks if str(x).strip())
                        else:
                            keys = tuple()
                        if gk:
                            gl_key_to_clause_keys[gk] = keys
                            gl_key_to_priority[gk] = pr
                except Exception:
                    clause_key_to_title = {}
                    gl_key_to_clause_keys = {}
                    gl_key_to_priority = {}
                try:
                    from collections import OrderedDict
                except Exception:
                    OrderedDict = dict  # type: ignore
                grouped_gl = OrderedDict()
                for _, row in df_gl.iterrows():
                    # Primary: map guideline key to clause_keys via template
                    gk = str(row.get("key", "")).strip()
                    keys_tuple = gl_key_to_clause_keys.get(gk)
                    keys: list[str]
                    if keys_tuple:
                        keys = list(keys_tuple)
                    else:
                        # Fallback: try to parse clause_keys column if present in CSV
                        cks_raw = row.get("clause_keys")
                        if isinstance(cks_raw, str) and cks_raw:
                            try:
                                import ast

                                parsed = ast.literal_eval(cks_raw)
                                if isinstance(parsed, (list, tuple)):
                                    keys = [str(x).strip() for x in parsed if str(x).strip()]
                                else:
                                    keys = []
                            except Exception:
                                keys = [s.strip() for s in cks_raw.split(",") if s.strip()]
                        else:
                            keys = []
                    heading = (
                        ", ".join([clause_key_to_title.get(k, k) for k in keys])
                        if keys
                        else "Unscoped"
                    )
                    grouped_gl.setdefault(heading, []).append(row)

                def _format_conf(val: Any) -> str:
                    try:
                        if val is None or (isinstance(val, float) and pd.isna(val)):
                            return ""
                        num = float(val)
                        if 0.0 <= num <= 1.0:
                            num *= 100.0
                        if num < 0:
                            num = 0.0
                        if num > 100:
                            num = 100.0
                        return f"{int(round(num))}%"
                    except Exception:
                        s = str(val).strip()
                        return f"{s}%" if s and not s.endswith("%") else s

                def _priority_class(p: str) -> str:
                    s = str(p or "").strip().lower()
                    if s in ("high", "material", "critical", "severe"):
                        return "sev-material"
                    if s in ("medium", "important", "major"):
                        return "sev-important"
                    if s in ("low", "minor"):
                        return "sev-minor"
                    return "sev-minor"

                for heading, rows in grouped_gl.items():
                    st.markdown(f"### {heading}")
                    for r in rows:
                        guideline_text = str(r.get("guideline", "")).strip()
                        gk = str(r.get("key", "")).strip()
                        prio = gl_key_to_priority.get(gk, "")
                        matched = str(r.get("guideline_matched", "")).strip()
                        conf_str = _format_conf(r.get("confidence", ""))
                        explanation = str(r.get("explanation", "")).strip()

                        # Badge
                        is_matched = str(matched).strip().lower() in ("true", "yes", "1")
                        if is_matched:
                            badge_html = '<span class="sev verified">Verified</span>'
                        else:
                            pr_cls = _priority_class(prio)
                            pr_label = (str(prio).strip() or "priority").capitalize()
                            badge_html = f'<span class="sev {pr_cls}">{pr_label}</span>'

                        conf_html = (
                            f'<span class="dp-conf">{conf_str}</span>'
                            if str(conf_str).strip()
                            else ""
                        )
                        expl_html = (
                            f'<span class="dp-meta">{explanation}</span>'
                            if str(explanation).strip()
                            else ""
                        )
                        meta_parts = [x for x in [conf_html, expl_html] if x]
                        meta_html = (
                            badge_html + (" " + " ".join(meta_parts) if meta_parts else "")
                        ).strip()

                        st.markdown(f"**{guideline_text}**")
                        st.markdown(meta_html, unsafe_allow_html=True)

    with tab_agent:
        repo_root = get_repo_root()
        src_dir = repo_root / "src"
        if str(src_dir) not in sys.path:
            sys.path.insert(0, str(src_dir))
        try:
            from contract_ai_core.contract_agent import Agent, AgentConfig
            from contract_ai_core.schema import (
                ClassifiedParagraph,
                ContractTypeTemplate,
                DocumentAnalysis,
                DocumentClassification,
                ExtractedDatapoint,
                ExtractionResult,
                Paragraph,
                ReviewedGuideline,
            )
        except Exception as e:
            st.error(f"Unable to load core library: {e}")
        else:
            # Load template
            try:
                tmpl_dict = load_template_dict(template_key)
                contract_type = ContractTypeTemplate.model_validate(tmpl_dict)
            except Exception as e:
                contract_type = None
                st.warning(f"Template load failed: {e}")

            # Helpers
            def _parse_conf_pct_to_unit(v: Any) -> float | None:
                try:
                    if v is None or (isinstance(v, float) and pd.isna(v)):
                        return None
                    s = str(v).strip()
                    if s == "":
                        return None
                    num = float(s)
                    if num > 1.0:
                        num = num / 100.0
                    return max(0.0, min(1.0, num))
                except Exception:
                    return None

            def _parse_evidence_list(s: Any) -> list[int] | None:
                if s is None:
                    return None
                try:
                    txt = str(s).strip()
                    if not txt:
                        return None
                    import ast

                    parsed = ast.literal_eval(txt)
                    if isinstance(parsed, (list, tuple)):
                        out: list[int] = []
                        for x in parsed:
                            try:
                                out.append(int(x))
                            except Exception:
                                continue
                        return out
                    if isinstance(parsed, (int, float)):
                        return [int(parsed)]
                except Exception:
                    pass
                return None

            # Build artifacts from CSVs
            df_cls = (
                load_classification_csv(template_key, model_name, current_path.stem)
                if model_name
                else None
            )
            if df_cls is not None and not df_cls.empty:
                cls_pars: list[ClassifiedParagraph] = []
                clause_to_pars: dict[str, list[int]] = {}
                for _, row in df_cls.iterrows():
                    try:
                        idx = int(str(row.get("index", "")).strip())
                    except Exception:
                        continue
                    text_i = str(row.get("text", "")).strip()
                    ck = str(row.get("clause_key", "")).strip() or None
                    conf_unit = _parse_conf_pct_to_unit(row.get("confidence", ""))
                    cp = ClassifiedParagraph(
                        paragraph=Paragraph(index=idx, text=text_i),
                        clause_key=ck,
                        confidence=conf_unit,
                    )
                    cls_pars.append(cp)
                    if ck:
                        clause_to_pars.setdefault(ck, []).append(idx)
                classification = DocumentClassification(
                    paragraphs=cls_pars, clause_to_paragraphs=clause_to_pars or None
                )
            else:
                classification = None

            df_dp = (
                load_datapoints_csv(template_key, model_name, current_path.stem)
                if model_name
                else None
            )
            if df_dp is not None and not df_dp.empty:
                dps: list[ExtractedDatapoint] = []
                for _, r in df_dp.iterrows():
                    key = str(r.get("key", "")).strip()
                    value = r.get("value", None)
                    expl = str(r.get("explanation", "")).strip() or None
                    ev = _parse_evidence_list(r.get("evidence", None))
                    conf = _parse_conf_pct_to_unit(r.get("confidence", ""))
                    dps.append(
                        ExtractedDatapoint(
                            key=key,
                            value=value,
                            explanation=expl,
                            evidence_paragraph_indices=ev,
                            confidence=conf,
                        )
                    )
                extraction = ExtractionResult(datapoints=dps)
            else:
                extraction = None

            df_gl = (
                load_guidelines_csv(template_key, model_name, current_path.stem)
                if model_name
                else None
            )
            if df_gl is not None and not df_gl.empty:
                gls: list[ReviewedGuideline] = []
                for _, r in df_gl.iterrows():
                    key = str(r.get("key", "")).strip()
                    matched_raw = str(r.get("guideline_matched", "")).strip().lower()
                    matched = True if matched_raw in ("true", "1", "yes", "y") else False
                    conf = _parse_conf_pct_to_unit(r.get("confidence", ""))
                    expl = str(r.get("explanation", "")).strip() or None
                    ev = _parse_evidence_list(r.get("evidence", None))
                    gls.append(
                        ReviewedGuideline(
                            key=key,
                            guideline_matched=matched,
                            confidence=conf,
                            explanation=expl,
                            evidence_paragraph_indices=ev,
                        )
                    )
            else:
                gls = []

            analysis = DocumentAnalysis(
                metadata=None,
                classified_clauses=classification,
                extracted_datapoints=extraction,
                reviewed_guidelines=gls or None,
            )

            agent_key = f"agent::{template_key}::{model_name}::{current_path.name}"
            if agent_key not in st.session_state:
                st.session_state[agent_key] = Agent(
                    contract_type=contract_type, analysis=analysis, config=AgentConfig()
                )
            agent: Agent = st.session_state[agent_key]

            # Render chat history
            if getattr(agent, "_history", None):
                for uq, aa in agent._history:
                    with st.chat_message("user"):
                        st.markdown(uq)
                    with st.chat_message("assistant"):
                        st.markdown(aa)

            # Chat input at bottom; on submit, process and rerun to keep input anchored
            user_q = st.chat_input("Ask a question about this contract…")
            if user_q and user_q.strip():
                try:
                    with st.spinner("Thinking …"):
                        agent.ask(user_q)
                except Exception as e:
                    st.error(f"Agent error: {e}")
                else:
                    st.rerun()

            # Reset button after the input
            if st.button("Reset conversation"):
                agent.reset()
                st.success("Conversation reset.")
                st.rerun()

    with tab_compare:
        # Compare current document to another from the same template folder or a template
        repo_root = get_repo_root()
        src_dir = repo_root / "src"
        if str(src_dir) not in sys.path:
            sys.path.insert(0, str(src_dir))
        # Imports
        try:
            from contract_ai_core.compare import DocumentCompare, DocumentCompareConfig
            from contract_ai_core.schema import (
                ClassifiedParagraph,
                DocumentClassification,
                Paragraph,
            )
            from contract_ai_core.utilities import text_to_paragraphs  # type: ignore
        except Exception as e:
            st.error(f"Unable to load compare module: {e}")
        else:
            # Add switch to choose between Contract and Template
            compare_type = st.radio(
                "Compare with:",
                ["Contract", "Template"],
                horizontal=True,
                key=f"compare_type::{template_key}::{current_path.name}",
            )

            if compare_type == "Template":
                # Find template files
                templates_dir = repo_root / "dataset" / "documents" / "templates" / template_key
                template_files = []
                if templates_dir.exists():
                    template_files = sorted(templates_dir.glob("*.docx"))

                if not template_files:
                    st.info(f"No template files found in {templates_dir.relative_to(repo_root)}.")
                    st.button("Compare", disabled=True)
                    st.stop()
                else:
                    template_names = [f.name for f in template_files]
                    state_key_template = f"compare_template::{template_key}::{current_path.name}"
                    current_template_val = st.session_state.get(state_key_template, "")

                    picked_template = st.selectbox(
                        "Select template",
                        [""] + template_names,
                        index=([""] + template_names).index(current_template_val)
                        if current_template_val in ([""] + template_names)
                        else 0,
                        help="Select a template to compare with.",
                    )

                    if picked_template != current_template_val:
                        # Clear any cached result for old selection
                        for k in list(st.session_state.keys()):
                            if k.startswith(
                                f"compare::{template_key}::{model_name}::{current_path.name}::template::"
                            ):
                                st.session_state.pop(k, None)
                    st.session_state[state_key_template] = picked_template

                    if not picked_template:
                        st.info("Select a template to compare.")
                        st.button("Compare", disabled=True)
                        st.stop()

                    target_path = next(f for f in template_files if f.name == picked_template)
                    is_template_comparison = True
            else:
                # File picker for the second document (contract)
                other_files = [p for p in files if p != current_path]
                if not other_files:
                    st.info("No other files found to compare against.")
                    st.button("Compare", disabled=True)
                    st.stop()
                else:
                    # Generate display names for other files using metadata
                    other_file_indices = [i for i, p in enumerate(files) if p != current_path]
                    other_display_names = [
                        get_contract_display_name(files[i], i) for i in other_file_indices
                    ]
                    other_names = [p.name for p in other_files]

                    # Map display names back to file names
                    display_to_filename = dict(zip(other_display_names, other_names, strict=False))

                    state_key_target = (
                        f"compare_target::{template_key}::{model_name}::{current_path.name}"
                    )
                    current_val = st.session_state.get(state_key_target, "")

                    # Find the index for the current selection
                    try:
                        current_display_idx = (
                            ([""] + other_names).index(current_val)
                            if current_val in ([""] + other_names)
                            else 0
                        )
                    except ValueError:
                        current_display_idx = 0

                    picked_display = st.selectbox(
                        "Compare current file with",
                        [""] + other_display_names,
                        index=current_display_idx,
                        help="Select another document to start comparison.",
                    )

                    # Convert display name back to filename
                    picked = display_to_filename.get(picked_display, "") if picked_display else ""

                    if picked != current_val:
                        # Clear any cached result for old selection
                        for k in list(st.session_state.keys()):
                            if k.startswith(
                                f"compare::{template_key}::{model_name}::{current_path.name}::"
                            ):
                                st.session_state.pop(k, None)
                    st.session_state[state_key_target] = picked
                    if not picked:
                        st.info("Select a document to compare.")
                        st.button("Compare", disabled=True)
                        st.stop()
                    target_path = next(p for p in other_files if p.name == picked)
                    is_template_comparison = False

            # Generate cache key based on comparison type
            compare_type_key = "template" if is_template_comparison else "contract"
            compare_cache_key = f"compare::{template_key}::{model_name}::{current_path.name}::{compare_type_key}::{target_path.name}"
            result = st.session_state.get(compare_cache_key)
            compare_clicked = st.button("Compare")
            if result is None and compare_clicked:
                with st.spinner("Comparing, please wait…"):
                    # Read and split paragraphs
                    try:
                        text1 = read_text_best_effort(current_path)
                        # Read second document based on type
                        if is_template_comparison:
                            text2 = read_docx_text(target_path)
                        else:
                            text2 = read_text_best_effort(target_path)
                        paras1 = [
                            Paragraph(index=i, text=pp.text)
                            for i, pp in enumerate(text_to_paragraphs(text1))
                        ]
                        paras2 = [
                            Paragraph(index=i, text=pp.text)
                            for i, pp in enumerate(text_to_paragraphs(text2))
                        ]

                    except Exception as e:
                        st.error(f"Failed to load documents: {e}")
                        paras1, paras2 = [], []

                    # Build classification for doc1 (for clause labels)
                    df_cls = (
                        load_classification_csv(template_key, model_name, current_path.stem)
                        if model_name
                        else None
                    )
                    classification: DocumentClassification | None
                    if df_cls is not None and not df_cls.empty:
                        cls_pars: list[ClassifiedParagraph] = []
                        clause_to_pars: dict[str, list[int]] = {}
                        for _, row in df_cls.iterrows():
                            try:
                                idx_i = int(str(row.get("index", "")).strip())
                            except Exception:
                                continue
                            txt = str(row.get("text", "")).strip()
                            ck = str(row.get("clause_key", "")).strip() or None
                            cp = ClassifiedParagraph(
                                paragraph=Paragraph(index=idx_i, text=txt),
                                clause_key=ck,
                                confidence=None,
                            )
                            cls_pars.append(cp)
                            if ck:
                                clause_to_pars.setdefault(ck, []).append(idx_i)
                        classification = DocumentClassification(
                            paragraphs=cls_pars, clause_to_paragraphs=clause_to_pars or None
                        )
                    else:
                        classification = None

                    # Run compare and cache result
                    if paras1 and paras2:
                        try:
                            comparer = DocumentCompare(
                                paragraphs_doc1=paras1,
                                paragraphs_doc2=paras2,
                                classification_doc1=classification,
                                config=DocumentCompareConfig(),
                            )
                            result = comparer.compare()
                            st.session_state[compare_cache_key] = result
                        except Exception as e:
                            import traceback

                            st.error(f"Compare failed: {e}" + "\n" + traceback.format_exc())
                            result = None

                if result is None and not compare_clicked:
                    st.info("Click Compare to run comparison.")
                elif not result:
                    st.info("No comparison result to display.")
                else:
                    # Summary of differences by severity
                    try:
                        items = getattr(result, "items", []) or []
                        cnt_material = sum(
                            1
                            for it in items
                            if str(getattr(it, "severity", "")).strip().lower() == "material"
                        )
                        cnt_important = sum(
                            1
                            for it in items
                            if str(getattr(it, "severity", "")).strip().lower() == "important"
                        )
                        cnt_minor = sum(
                            1
                            for it in items
                            if str(getattr(it, "severity", "")).strip().lower() == "minor"
                        )
                        st.markdown("### Summary of differences by severity")
                        st.markdown(f"**Material:** {cnt_material}")
                        st.markdown(f"**Important:** {cnt_important}")
                        st.markdown(f"**Minor:** {cnt_minor}")
                    except Exception:
                        pass
                    # Build clause key to title mapping
                    try:
                        tmpl = load_template_dict(template_key)
                        clauses = tmpl.get("clauses", []) or []
                        clause_title_by_key = {
                            str(c.get("key")): (c.get("title") or str(c.get("key")))
                            for c in clauses
                        }
                    except Exception:
                        clause_title_by_key = {}

                    para_cells: list[str] = []
                    clause_cells: list[str] = []
                    expl_cells: list[str] = []
                    prev_ck: str | None = None
                    for it in result.items:
                        ck = str(getattr(it, "clause_key_doc1", "") or "").strip()
                        is_repeat = prev_ck is not None and ck != "" and ck == prev_ck
                        if ck != "":
                            prev_ck = ck

                        # Clause cell (hide repeats)
                        clause_title = clause_title_by_key.get(ck, ck)
                        if is_repeat:
                            clause_cells.append('<div class="repeat-cell"></div>')
                        else:
                            clause_cells.append(f'<div class="normal-cell">{clause_title}</div>')

                        # Paragraph cell (add right border when repeat)
                        paragraph_html = getattr(it, "text_markup", "") or ""
                        if is_repeat:
                            para_cells.append(
                                f'<div class="text-repeat-cell">{paragraph_html}</div>'
                            )
                        else:
                            para_cells.append(
                                f'<div class="text-normal-cell">{paragraph_html}</div>'
                            )

                        # Explanation cell (severity + confidence + rationale)
                        severity = getattr(it, "severity", "") or ""
                        rationale = getattr(it, "rationale", "") or ""
                        conf_val = getattr(it, "confidence", None)
                        try:
                            conf_str = (
                                f"{int(round(float(conf_val) * 100))}%"
                                if conf_val is not None
                                else ""
                            )
                        except Exception:
                            conf_str = ""
                        parts = []
                        if severity:
                            parts.append(
                                f'<span class="sev sev-{severity.lower()}">{severity}</span>'
                            )
                        if conf_str:
                            parts.append(f'<span class="cmp-conf">{conf_str}</span>')
                        if rationale:
                            parts.append(rationale)
                        expl_cells.append(" ".join(parts).strip())

                    df_view = pd.DataFrame(
                        {
                            "paragraph": pd.Series(para_cells),
                            "clause": pd.Series(clause_cells),
                            "explanation": pd.Series(expl_cells),
                        }
                    )
                    html = df_view.to_html(index=False, escape=False)
                    html = html.replace(
                        '<table border="1" class="dataframe">',
                        (
                            "<style>"
                            ".cmp-table{width:100%;table-layout:fixed;border-collapse:collapse;}"
                            ".cmp-table th,.cmp-table td{border:none;padding:0.5rem;vertical-align:top;word-wrap:break-word;white-space:normal;}"
                            ".cmp-table col.diff-col{width:60%;}"
                            ".cmp-table col.clause-col{width:15%;}"
                            ".cmp-table col.expl-col{width:25%;}"
                            ".cmp-table td .text-repeat-cell{border-right:4px solid #888;padding-right:0.5rem;}"
                            ".diff-block{background:#f9f9f9;padding:0.5rem;border-radius:4px;}"
                            ".h-add{color:#1b5e20;display:block;}"
                            ".h-del{color:#b71c1c;display:block;}"
                            ".h-eq{color:#555;display:block;}"
                            ".h-meta{color:#888;display:block;font-style:italic;}"
                            ".sev{padding:2px 6px;border-radius:10px;font-size:0.85em;color:#fff;}"
                            ".cmp-conf{color:#2e7d32;font-weight:600;}"
                            ".sev-material{background:#b71c1c;}"
                            ".sev-important{background:#ef6c00;}"
                            ".sev-minor{background:#616161;}"
                            ".sev-none{background:#2e7d32;}"
                            "</style>"
                            '<table class="cmp-table">'
                            "<colgroup>"
                            '<col class="diff-col"/>'
                            '<col class="clause-col"/>'
                            '<col class="expl-col"/>'
                            "</colgroup>"
                        ),
                    )
                    st.markdown(html, unsafe_allow_html=True)


if __name__ == "__main__":
    main()
