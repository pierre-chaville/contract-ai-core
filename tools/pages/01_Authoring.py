from __future__ import annotations

import os
import sys
from pathlib import Path
from typing import Any, Optional

import streamlit as st


def get_repo_root() -> Path:
    here = Path(__file__).resolve()
    dataset_dir: Optional[Path] = None
    for p in here.parents:
        if p.name == "dataset":
            dataset_dir = p
            break
    return dataset_dir.parent if dataset_dir else here.parents[2]


def list_contract_types() -> list[str]:
    repo_root = get_repo_root()
    ct_dir = repo_root / "dataset" / "contract_types"
    keys: list[str] = []
    try:
        for p in sorted(ct_dir.glob("*.json")):
            keys.append(p.stem)
    except Exception:
        pass
    return keys


def list_authoring_templates(contract_type: str) -> list[Path]:
    """Find available authoring templates (.docx, .txt, .md)."""
    repo_root = get_repo_root()
    candidates = [repo_root / "dataset" / "documents" / "templates" / contract_type]
    files: list[Path] = []
    for base in candidates:
        try:
            if base.exists():
                files.extend(sorted(base.glob("*.docx")))
                files.extend(sorted(base.glob("*.txt")))
                files.extend(sorted(base.glob("*.md")))
        except Exception:
            continue
    return files


def read_text_best_effort(path: Path) -> str:
    encodings = [
        "utf-8",
        "utf-8-sig",
        "cp1252",
        "latin-1",
        "utf-16",
        "utf-16-le",
        "utf-16-be",
    ]
    for enc in encodings:
        try:
            return path.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def extract_paragraphs_from_docx(path: Path) -> list[str]:
    try:
        import docx  # type: ignore
    except Exception as e:
        raise RuntimeError(
            "python-docx is required to read .docx templates. Install with: pip install python-docx"
        ) from e
    doc = docx.Document(str(path))
    return [p.text for p in doc.paragraphs]


def html_escape(s: str) -> str:
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def render_docx_to_html(path: Path) -> str:
    try:
        import mammoth  # type: ignore

        with open(path, "rb") as f:
            result = mammoth.convert_to_html(f)
        return result.value  # type: ignore[attr-defined]
    except Exception:
        try:
            import docx  # type: ignore

            d = docx.Document(str(path))
            parts = [f"<p>{html_escape(p.text)}</p>" for p in d.paragraphs]
            return "\n".join(parts)
        except Exception:
            return "<p>(Unable to preview .docx content)</p>"


def render_text_to_html(text: str) -> str:
    if not text:
        return ""
    ch = html_escape(text)
    paras = [p.strip() for p in ch.split("\n\n")]
    return "\n".join(f"<p>{p}</p>" for p in paras if p)


def main() -> None:
    st.set_page_config(page_title="Authoring", layout="wide")
    st.title("Authoring")

    repo_root = get_repo_root()
    src_dir = repo_root / "src"
    if str(src_dir) not in sys.path:
        sys.path.insert(0, str(src_dir))

    # Lazy imports after sys.path update
    try:
        from contract_ai_core.authoring import AuthoringConfig, AuthoringDocument  # type: ignore
    except Exception as e:
        st.error(f"Unable to load authoring module: {e}")
        st.stop()

    tab_author, tab_template, tab_variables, tab_draft = st.tabs(
        ["Author", "Template", "Variables", "Draft"]
    )

    with tab_author:
        ct_keys = list_contract_types()

        col1, col2 = st.columns([1, 2])
        with col1:
            contract_type = st.selectbox(
                "Contract type",
                ct_keys if ct_keys else ["(none found)"],
                index=0 if ct_keys else 0,
            )
            templates = list_authoring_templates(contract_type)
            template_names = ["(none)"] + [os.path.basename(p) for p in templates]
            picked_template = st.selectbox(
                "Template file",
                template_names,
                index=0,
                help="Pick a .docx/.txt/.md template containing docxtpl placeholders.",
            )
        with col2:
            context_text = st.text_area(
                "Context / Instructions",
                height=220,
                placeholder=(
                    "Explain how to fill placeholders, enumerations, conditions, and loops.\n"
                    "E.g., Party A is ACME Inc.; Effective Date is the signature date; fees per tier..."
                ),
            )

        if st.button("Apply"):
            # Build template paragraphs from selection
            paragraphs: list[str] = []
            template_text_capture: str | None = None
            if picked_template != "(none)":
                abs_path = (
                    repo_root
                    / "dataset"
                    / "documents"
                    / "templates"
                    / contract_type
                    / picked_template
                )
                try:
                    if abs_path.suffix.lower() == ".docx":
                        paragraphs = extract_paragraphs_from_docx(abs_path)
                        template_text_capture = "\n".join(paragraphs)
                    else:
                        text = read_text_best_effort(abs_path)
                        paragraphs = [
                            p.strip() for p in text.replace("\r\n", "\n").split("\n\n") if p.strip()
                        ]
                        template_text_capture = text
                except Exception as e:
                    st.error(f"Failed to read template: {e}")
                    paragraphs = []
                    template_text_capture = None

            if not paragraphs:
                st.warning(
                    "No template content provided or template is empty. Using an empty template excerpt."
                )

            try:
                authoring = AuthoringDocument(
                    template_paragraphs=paragraphs,
                    context_text=context_text,
                    config=AuthoringConfig(),
                )
                output = authoring.generate_fields()
            except Exception as e:
                st.error(f"Authoring failed: {e}")
                output = None

            st.session_state["authoring_output"] = output
            # store template preview metadata
            if picked_template != "(none)":
                st.session_state["authoring_template_path"] = str(abs_path)
                if template_text_capture is None:
                    try:
                        # Fallback join
                        st.session_state["authoring_template_text"] = "\n\n".join(paragraphs)
                    except Exception:
                        st.session_state["authoring_template_text"] = ""
                else:
                    st.session_state["authoring_template_text"] = template_text_capture
            # Attempt to render DOCX draft via docxtpl if a .docx template is selected
            if picked_template != "(none)":
                abs_path = (
                    repo_root
                    / "dataset"
                    / "documents"
                    / "templates"
                    / contract_type
                    / picked_template
                )
                if abs_path.suffix.lower() == ".docx" and output is not None:
                    try:
                        import json

                        try:
                            from docxtpl import DocxTemplate  # type: ignore
                        except Exception as e:
                            raise RuntimeError(
                                "docxtpl is required to render .docx drafts. Install with: pip install docxtpl"
                            ) from e
                        # Build context mapping from output fields
                        mapping: dict[str, Any] = {}
                        for f in getattr(output, "fields", []) or []:
                            key = getattr(f, "key", None)
                            if not key:
                                continue
                            vjson = getattr(f, "value_json", None)
                            if vjson is not None and str(vjson).strip() != "":
                                try:
                                    mapping[key] = json.loads(str(vjson))
                                except Exception:
                                    mapping[key] = vjson
                            else:
                                mapping[key] = getattr(f, "value", None)

                        # Render
                        tpl = DocxTemplate(str(abs_path))
                        tpl.render(mapping)
                        out_dir = repo_root / "tools" / "output" / "drafts" / contract_type
                        out_dir.mkdir(parents=True, exist_ok=True)
                        out_path = out_dir / (Path(picked_template).stem + "_draft.docx")
                        tpl.save(str(out_path))
                        st.session_state["authoring_draft_path"] = str(out_path)

                        # Extract text for display
                        try:
                            import docx  # type: ignore

                            d = docx.Document(str(out_path))
                            draft_text = "\n".join(p.text for p in d.paragraphs)
                        except Exception:
                            draft_text = "(Unable to preview .docx content)"
                        st.session_state["authoring_draft_text"] = draft_text
                        st.success("Draft generated. See the 'Draft' tab.")
                    except Exception as e:
                        st.warning(f"Draft generation skipped: {e}")
                else:
                    st.info("Select a .docx template to generate a draft document.")
            else:
                st.success("Variables extracted. See the 'Variables' tab.")

    with tab_variables:
        out = st.session_state.get("authoring_output")
        if not out:
            st.info("Run authoring in the Author tab to see extracted variables.")
        else:
            try:
                fields = getattr(out, "fields", [])
            except Exception:
                fields = []
            if not fields:
                st.info("No variables extracted.")
            else:
                for f in fields:
                    key = getattr(f, "key", "")
                    val = getattr(f, "value", None)
                    conf = getattr(f, "confidence", None)
                    expl = getattr(f, "explanation", "") or ""
                    # Render
                    st.markdown(f"**{key}**")
                    st.code(val)
                    conf_str = ""
                    try:
                        conf_str = f"{float(conf):.2f}" if conf is not None else ""
                    except Exception:
                        conf_str = str(conf) if conf is not None else ""
                    meta = " ".join([p for p in [conf_str, expl] if p]).strip()
                    if meta:
                        st.caption(meta)

    with tab_template:
        tpath = st.session_state.get("authoring_template_path")
        ttext = st.session_state.get("authoring_template_text")
        if not tpath:
            st.info("Select a template and click Apply to preview the template content.")
        else:
            st.markdown(f"Template: `{tpath}`")
            abs_path = Path(str(tpath))
            if abs_path.suffix.lower() == ".docx":
                html = render_docx_to_html(abs_path)
            else:
                html = render_text_to_html(ttext or "")
            st.markdown(html, unsafe_allow_html=True)

    with tab_draft:
        draft_path = st.session_state.get("authoring_draft_path")
        draft_text = st.session_state.get("authoring_draft_text")
        if not draft_path:
            st.info(
                "No draft generated yet. Use the Author tab to apply and generate a draft from a .docx template."
            )
        else:
            st.subheader("Draft preview")
            try:
                with open(draft_path, "rb") as f:
                    data = f.read()
                st.download_button(
                    label="Export DOCX",
                    data=data,
                    file_name=Path(draft_path).name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            except Exception as e:
                st.warning(f"Unable to load generated file for download: {e}")
            try:
                html = render_docx_to_html(Path(str(draft_path)))
            except Exception:
                html = render_text_to_html(draft_text or "")
            st.markdown(html, unsafe_allow_html=True)


if __name__ == "__main__":
    main()
